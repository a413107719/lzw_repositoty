import os
import shutil
import docx
from docx import Document
import openpyxl


def main():
    global new_mediapath
    new_mediapath = input_folder + '\\media'
    if os.path.isdir(new_mediapath):
        shutil.rmtree(new_mediapath)
        print('删除media文件夹')
    # 找出所有的docx文件
    documents = os.listdir(input_folder)
    docs = []
    for document in documents:
        if document.split('.')[1] == 'docx':
            docs.append(document)
    print('文件夹中包含word文档：' + str(docs))

    allpointinformation = []
    for doc in docs:
        docname = doc.split('.')[0]
        docpath = input_folder + "\\" + doc
        document = Document(docpath)  # 打开.docx文件
        find_allmedia(docname, docpath)
        find_all_table(document,allpointinformation)

    for i in allpointinformation:
        print(i)
    write2excl(allpointinformation)


def write2excl(list):
    wb = openpyxl.Workbook()
    sheet = wb.active  # 选择一个excel页面
    # 表格填入数据
    sheet.cell(row=1, column=1).value = "编号"
    sheet.cell(row=1, column=2).value = "名称"
    sheet.cell(row=1, column=3).value = "类型"
    sheet.cell(row=1, column=4).value = "细类"
    sheet.cell(row=1, column=5).value = "年代"
    sheet.cell(row=1, column=6).value = "统计年代"
    sheet.cell(row=1, column=7).value = "级别"
    sheet.cell(row=1, column=8).value = "用途"
    sheet.cell(row=1, column=9).value = "地址"
    sheet.cell(row=1, column=10).value = "经度"
    sheet.cell(row=1, column=11).value = "纬度"
    sheet.cell(row=1, column=12).value = "海拔"
    sheet.cell(row=1, column=13).value = "测点说明"
    sheet.cell(row=1, column=14).value = "数量"
    sheet.cell(row=1, column=15).value = "说明"
    sheet.cell(row=1, column=16).value = "面积"
    sheet.cell(row=1, column=17).value = "所有权"
    sheet.cell(row=1, column=18).value = "使用单位"
    sheet.cell(row=1, column=19).value = "隶属"
    sheet.cell(row=1, column=20).value = "现状评估"
    sheet.cell(row=1, column=21).value = "现状描述"
    sheet.cell(row=1, column=22).value = "损毁原因_自然因素"
    sheet.cell(row=1, column=23).value = "损毁原因_人为因素"
    sheet.cell(row=1, column=24).value = "损毁原因_原因描述"
    sheet.cell(row=1, column=25).value = "自然环境"
    sheet.cell(row=1, column=26).value = "人文环境"
    sheet.cell(row=1, column=27).value = "普查组建议"
    sheet.cell(row=1, column=28).value = "简介"
    sheet.cell(row=1, column=29).value = "照片说明"
    sheet.cell(row=1, column=30).value = "照片拍摄时间"

    for rownum, rowlist in enumerate(list):
        for colnum, text in enumerate(rowlist):
            sheet.cell(row=rownum + 2, column=colnum+1).value = text

    wb.save(input_folder + os.sep + '文物信息提取.xlsx')
    print('已经下载数据')


# 找出word内的所有table
def find_all_table(document,allinfor):
    # 第2张表
    table = document.tables[1]
    table_imfor = extract_excel_imformation(table)
    # # 打印原始表
    # for num, row in enumerate(table_imfor):
    #     print(num, row)
    # print()
    # 取值
    mingchen = table_imfor[0][1]
    dizhi = table_imfor[1][1]
    for i in range(5, 11):
        if "●" in table_imfor[i][1]:
            leixing = table_imfor[i][1].split("●")[1]
            xilei = table_imfor[i][3].split("●")[1].split("〇")[0]
            break
    niandai = table_imfor[11][1]
    tongjiniandai = table_imfor[12][3].split("■")[1].split("□")[0]
    mianji = table_imfor[13][1]
    suoyouquan = table_imfor[14][3].split("■")[1].split("□")[0]
    shiyongdanwei = table_imfor[15][5]
    nishu = table_imfor[15][9]
    yongtu = get_multi_value(table_imfor[16][3])
    jibie = table_imfor[17][3].split("●")[1].split("〇")[0]

    # 第3张表
    table = document.tables[2]
    table_imfor = extract_excel_imformation(table)
    # # 打印原始表
    # for num, row in enumerate(table_imfor):
    #     print(num, row)
    # print()
    # 取值
    shuliang  = table_imfor[0][2]
    shuoming = table_imfor[1][2]
    jianjie = table_imfor[2][2]
    xzpg = table_imfor[3][3].split("●")[1].split("〇")[0]
    xzms = table_imfor[4][3]
    shyy_zrys = get_multi_value(table_imfor[5][3])
    shyy_rwys = get_multi_value(table_imfor[6][3])
    shyy_ms = table_imfor[7][3]

    # 第4张表
    table = document.tables[3]
    table_imfor = extract_excel_imformation(table)
    # # 打印原始表
    # for num, row in enumerate(table_imfor):
    #     print(num, row)
    # print()
    # 取值
    zrhj = table_imfor[0][2]
    rwhj = table_imfor[1][2]
    pczjy = table_imfor[2][2]

    # 第5张表
    table = document.tables[4]
    table_imfor = extract_excel_imformation(table)
    # # 打印原始表
    # for num, row in enumerate(table_imfor):
    #     print(num, row)
    # print()
    # 取值
    bh = table_imfor[2][0]
    jd = get_xy(table_imfor[2][1])
    wd = get_xy(table_imfor[2][2])
    hb = table_imfor[2][3]
    cdsm = table_imfor[2][4]

    # 第9张表
    table = document.tables[8]
    table_imfor = extract_excel_imformation(table)
    # 打印原始表
    # for num, row in enumerate(table_imfor):
    #     print(num, row)
    # print()
    # 取值
    zpsm = table_imfor[2][1]
    pssj = table_imfor[1][3]

    # print(zpsm,pssj)
    # print(mingchen,dizhi,cedianshuoming)
    # print(leixing,xilei,niandai,tongjiniandai,mianji,suoyouquan)
    # print(shiyongdanwei,nishu,yongtu,jibie)
    # print(shuliang,shuoming,jianjie)
    # print(xzpg,xzms,shyy_zrys,shyy_rwys,shyy_ms)
    # print(zrhj,rwhj,pczjy)
    # print(bh,jd,wd,hb,cdsm)
    # print()

    allinfor.append([
        bh, mingchen, leixing, xilei, niandai, tongjiniandai, jibie,
        yongtu, dizhi, jd, wd, hb, cdsm, shuliang, shuoming, mianji,
        suoyouquan, shiyongdanwei, nishu, xzpg, xzms, shyy_zrys,
        shyy_rwys, shyy_ms, zrhj, rwhj, pczjy,  jianjie, zpsm, pssj
    ])


def get_xy(text):
    du = text.split('°')[0]
    fen = text.split('°')[1].split('\'')[0]
    miao =text.split('\'')[1].split('"')[0]
    # newtext = du + fen/60+miao/3600
    print(du,fen,miao)
    newtext= float(du) + float(fen)/60 + float(miao)/3600
    print(newtext)
    return newtext


def get_multi_value(oldvalue):
    list = []
    for i in oldvalue.split("■"):
        y = i.split('□')[0]
        if y != '':
            list.append(y)
    newvalue = ",".join(list)
    return newvalue


def extract_excel_imformation(table):
    tablerow = len(table.rows)
    tablecolumn = len(table.columns)
    table_imformation = []
    for i in range(0, tablerow):  # 从表格第二行开始循环读取表格数据
        rowvalue = []
        for y in range(0, tablecolumn):
            cellvalue = table.cell(i, y).text
            cellvalue = cellvalue.replace(' ', '')
            cellvalue = cellvalue.replace('\n', '')
            cellvalue = cellvalue.replace('\\', '')
            rowvalue.append(cellvalue)
        table_imformation.append(rowvalue)
    return table_imformation


# 找出word内的所有图片
def find_allmedia(docname, docpath):
    import zipfile
    f = zipfile.ZipFile(docpath)
    for medianame in f.namelist():
        if 'word/media/' in medianame:
            # print(medianame)
            f.extract(medianame, path=input_folder)
    # 整理图片信息
    if not os.path.isdir(new_mediapath):
        os.makedirs(new_mediapath)
        # print("成功创建media文件夹", '\n')
    mediapath = input_folder + '\\word\\media'
    for jpeg in os.listdir(mediapath):
        # print(jpeg)
        old_mediapath = mediapath + os.sep + jpeg
        shutil.copy2(old_mediapath, new_mediapath + os.sep + docname + '_' + jpeg)
    shutil.rmtree(input_folder + '\\word')
    # print('成功整理图片信息')


if __name__ == '__main__':
    input_folder = 'L:\\三普信息提取'
    main()

