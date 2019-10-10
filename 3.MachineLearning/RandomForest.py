from sklearn.metrics import confusion_matrix as con
from sklearn.metrics import zero_one_loss
from sklearn.ensemble import AdaBoostClassifier
import matplotlib.pyplot as plt
from collections import OrderedDict
from prettytable import PrettyTable
from sklearn.model_selection import cross_val_score
from sklearn.neighbors import KNeighborsClassifier
from pylab import mpl
import numpy as np
from sklearn.tree import DecisionTreeClassifier
import pydotplus
from sklearn import tree
from IPython.display import Image
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import make_scorer, accuracy_score
from sklearn.model_selection import GridSearchCV
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.inspection import plot_partial_dependence
import shutil
import os
import docx
from docx.shared import Inches


def prettytableprint(name, data):
    f = open(tempt_path + name + '.txt', 'wb')
    f.write(bytes(str(data), encoding='utf-8'))
    f.close()


# 打印标签统计表
def print_featurelable():
    print('--- 标签特征统计表 ---')
    prettytableprint('03_标签特征统计表1', '3.标签特征统计表')
    yvalue_table = data_train[str(data_train.columns[len(data_train.columns) - 1])].value_counts()
    table2 = PrettyTable(['值', '数量', '占比'])
    for i in range(len(yvalue_table)):
        value = yvalue_table.index[i]
        number = yvalue_table[value]
        table2.add_row([value, number, round(number / len(X), 2)])
    print(table2)
    prettytableprint('03_标签特征统计表2', table2)
    print()


# 不同n_estimators取值对应误差大小
def findbest_n_estimators(max_trees_num):
    learning_rate = 1

    dt_stump = DecisionTreeClassifier(max_depth=1, min_samples_leaf=1, max_features=2)
    dt_stump.fit(X_train, y_train)
    dt_stump_err = 1.0 - dt_stump.score(X_test, y_test)

    dt = DecisionTreeClassifier(max_depth=9, min_samples_leaf=1)
    dt.fit(X_train, y_train)
    dt_err = 1.0 - dt.score(X_test, y_test)

    ada_discrete = AdaBoostClassifier(
        base_estimator=dt_stump,
        learning_rate=learning_rate,
        n_estimators=max_trees_num,
        algorithm="SAMME")
    ada_discrete.fit(X_train, y_train)

    ada_real = AdaBoostClassifier(
        base_estimator=dt_stump,
        learning_rate=learning_rate,
        n_estimators=max_trees_num,
        algorithm="SAMME.R")
    ada_real.fit(X_train, y_train)

    fig = plt.figure()
    ax = fig.add_subplot(111)

    ax.plot([1, max_trees_num], [dt_stump_err] * 2, 'k-',
            label='Decision Stump Error')
    ax.plot([1, max_trees_num], [dt_err] * 2, 'k--',
            label='Decision Tree Error')

    ada_discrete_err = np.zeros((max_trees_num,))
    for i, y_pred in enumerate(ada_discrete.staged_predict(X_test)):
        ada_discrete_err[i] = zero_one_loss(y_pred, y_test)

    ada_discrete_err_train = np.zeros((max_trees_num,))
    for i, y_pred in enumerate(ada_discrete.staged_predict(X_train)):
        ada_discrete_err_train[i] = zero_one_loss(y_pred, y_train)

    ada_real_err = np.zeros((max_trees_num,))
    for i, y_pred in enumerate(ada_real.staged_predict(X_test)):
        ada_real_err[i] = zero_one_loss(y_pred, y_test)

    ada_real_err_train = np.zeros((max_trees_num,))
    for i, y_pred in enumerate(ada_real.staged_predict(X_train)):
        ada_real_err_train[i] = zero_one_loss(y_pred, y_train)

    ax.plot(np.arange(max_trees_num) + 1, ada_discrete_err,
            label='Discrete AdaBoost Test Error',
            color='red')
    ax.plot(np.arange(max_trees_num) + 1, ada_discrete_err_train,
            label='Discrete AdaBoost Train Error',
            color='blue')
    ax.plot(np.arange(max_trees_num) + 1, ada_real_err,
            label='Real AdaBoost Test Error',
            color='orange')
    ax.plot(np.arange(max_trees_num) + 1, ada_real_err_train,
            label='Real AdaBoost Train Error',
            color='green')

    ax.set_ylim((0, 0.5))
    ax.set_xlabel('n_estimators')
    ax.set_ylabel('error rate')

    leg = ax.legend(loc='upper right', fancybox=True)
    leg.get_frame().set_alpha(0.7)

    prettytableprint('05_不同n_estimators取值对应误差大小1', '5.不同n_estimators取值对应误差大小')
    plt.savefig(tempt_path + '05_不同n_estimators取值对应误差大小2' + '.png')
    plt.show()


# 不同max_features取值对应误差大小
def findbest_maxfeatures():
    # 据此修改https://scikit-learn.org/dev/auto_examples/ensemble/plot_ensemble_oob.html#sphx-glr-auto-examples-ensemble-plot-ensemble-oob-py
    RANDOM_STATE = 123
    ensemble_clfs = []
    x_columns = data_train.columns[1:len(data_train.columns) - 1]
    for i in range(len(x_columns)):
        ensemble_clfs.append((i + 1,
                              RandomForestClassifier(n_estimators=200,
                                                     warm_start=True, oob_score=True,
                                                     max_features=i + 1,
                                                     random_state=RANDOM_STATE)))

    error_rate = OrderedDict((label, []) for label, _ in ensemble_clfs)

    print('-- 不同max_features取值对应误差大小 --')
    prettytableprint("04_不同max_features取值对应误差大小1", '4.不同max_features取值对应误差大小')
    table3 = PrettyTable(['随机特征个数', '误差率'])
    for label, clf in ensemble_clfs:
        # for i in range(min_estimators, max_estimators + 1):
        #     clf.set_params(n_estimators=i)
        clf.fit(X, y)

        # Record the OOB error for each `n_estimators=i` setting.
        oob_error = 1 - clf.oob_score_
        error_rate[label].append((round(oob_error, 4)))
        # print(label)
        # print(round(oob_error,4))
        table3.add_row([label, round(oob_error, 4)])
    print(table3)
    prettytableprint("04_不同max_features取值对应误差大小2", table3)
    print()


# 建立模型
def build_model():
    print("开始建立模型！模型需要自动修正参数，请等待。。。。。")
    # 选择分类器的类型
    forest = RandomForestClassifier()
    # 可以通过定义树的各种参数，限制树的大小，防止出现过拟合现象
    # 模型需要运行两次，第一次树的数量默认值为10，第二次根据可视化图输入最合适值。[50, 100, 200, 300]
    parameters = {'n_estimators': [50],
                  'max_features': ['log2', 'sqrt', 'auto'],
                  'criterion': ['entropy', 'gini'],  # 分类标准用熵，基尼系数
                  'max_depth': [2, 3, 5, 10],
                  'min_samples_split': [2, 3, 5],
                  'min_samples_leaf': [1, 5, 8]

                  }
    # 以下是用于比较参数好坏的评分，使用'make_scorer'将'accuracy_score'转换为评分函数
    acc_scorer = make_scorer(accuracy_score)
    # 自动调参，GridSearchCV，它存在的意义就是自动调参，只要把参数输进去，就能给出最优化的结果和参数
    # GridSearchCV用于系统地遍历多种参数组合，通过交叉验证确定最佳效果参数。
    grid_obj = GridSearchCV(forest, parameters, scoring=acc_scorer, cv=5)
    grid_obj = grid_obj.fit(X_train, y_train)
    # 将forest设置为参数的最佳组合
    forest = grid_obj.best_estimator_
    print("建树参数:", forest, sep='\n')
    prettytableprint('06_建树参数1', "6.建树参数")
    prettytableprint('06_建树参数2', forest)
    # 将最佳算法运用于随机森林，随机森林建树
    forest.fit(X_train, y_train) # 用训练集数据训练模型
    print()
    return forest


# 单因子分析
def singlefeature(totalnum):
    prettytableprint('11_单因子分析0', '11_单因子分析')
    x = 0
    num = 0
    while x < totalnum:
        features = []

        i = x
        while i < totalnum and i-x < 2:
            features.append(i)
            i += 1
        plot_partial_dependence(forest, X_train, features, feature_names=feature_column, n_jobs=5, grid_resolution=20)
        fig = plt.gcf()
        fig.subplots_adjust(wspace=0.8, hspace=0.3)
        x += 2
        num += 1
        plt.savefig(tempt_path + '11_单因子分析' + str(num) + '.png')
    plt.show()


# 交叉验证
def cross_validation():
    k_range = range(1, 31)
    cv_scores = []		# 用来放每个模型的结果值
    for n in k_range:
        # knn模型，这里一个超参数可以做预测，当多个超参数时需要使用另一种方法GridSearchCV
        knn = KNeighborsClassifier(n)
        # cv：选择每次测试折数  accuracy：评价指标是准确度,可以省略使用默认值
        scores = cross_val_score(knn, X_train, y_train, scoring='accuracy', cv=5)
        cv_scores.append(scores.mean())
    max_score = max(cv_scores)
    max_index = cv_scores.index(max_score)

    # 绘出图形
    plt.plot(k_range, cv_scores)
    mpl.rcParams['font.sans-serif'] = ['SimHei']
    plt.title("交叉验证", fontsize=18)
    plt.xlabel('K_range')
    plt.ylabel('Accuracy')		# 通过图像选择最好的参数
    plt.savefig(tempt_path + '10_交叉验证3' + '.png')
    plt.show()
    best_knn = KNeighborsClassifier(n_neighbors=max_index)	 # 选择最优的K=3传入模型
    best_knn.fit(X_train, y_train)			# 训练模型
    print('------ 交叉验证 ------')
    table4 = PrettyTable(['最优K值', '最优评分'])
    table4.add_row([max_index, round(best_knn.score(X_test, y_test), 4)])
    print(table4)
    prettytableprint('10_交叉验证1', "10.交叉验证")
    prettytableprint('10_交叉验证2', table4)
    print()


# 打印混淆矩阵
def print_mxjuzhen(forest):
    # result = forest.fit(X_train, y_train).predict(X_test)
    result = forest.fit(X_train, y_train).predict(X_train)
    print('------------ 随机森林模型混淆矩阵 ------------')
    # table = con(y_test, result)
    table = con(y_train, result)
    table_x = PrettyTable(['实际类别', '预测适宜', '预测不适宜', '分类误差率(%)'])
    table_x.add_row(["适宜", table[0][0], table[0][1], round(table[0][1] / (table[0][1] + table[0][0]) * 100, 2)])
    table_x.add_row(["不适宜", table[1][0], table[1][1], round(table[1][0] / (table[1][1] + table[1][0]) * 100, 2)])
    print(table_x)
    prettytableprint('07_随机森林模型混淆矩阵1', "7.随机森林模型混淆矩阵")
    prettytableprint('07_随机森林模型混淆矩阵2', table_x)
    tablescore = (table[0][0] +table[1][1])/(table[0][0] +table[0][1] + table[1][0] +table[1][1])*100
    print('模型训练精度： ' + str(round(tablescore, 2)) + '%' + '\n')
    prettytableprint('07_随机森林模型混淆矩阵3', '模型训练精度： ' + str(round(tablescore, 2)) + '%')
    print()

    result1 = forest.fit(X_train, y_train).predict(X_test)
    # result = forest.fit(X_train, y_train).predict(X_train)
    print('--------- 测试数据预测结果 ---------')
    table = con(y_test, result1)
    # table = con(y_train, result)
    table_x = PrettyTable(['实际类别', '预测适宜', '预测不适宜'])
    table_x.add_row(["适宜", table[0][0], table[0][1]])
    table_x.add_row(["不适宜", table[1][0], table[1][1]])
    print(table_x)
    prettytableprint('08_测试数据预测结果1', "8.测试数据预测结果")
    prettytableprint('08_测试数据预测结果2', table_x)
    score = forest.score(X_test, y_test)  # 测试准确率
    print('模型准确率(泛化程度)： ' + str(round(score * 100, 2)) + '%' + '\n')
    prettytableprint('08_测试数据预测结果3', '模型准确率(泛化程度)： ' + str(round(score * 100, 2)) + '%')


# 因子重要性判断
def top_variable_importance(forest):
    feat_labels = data_train.columns[1:len(data_train.columns) - 1]
    # 对训练好的随机森林，完成重要性评估
    importances = forest.feature_importances_
    # print(importances)
    x_columns = data_train.columns[1:len(data_train.columns)-1]
    indices = np.argsort(importances)[::-1]
    # print(indices)
    np.argsort(x_columns,)
    list = []
    print('-- 所有变量影响程度排序 --')
    table3 = PrettyTable(['序号', '评价特征', '重要性(%)'])
    for f in range(X_train.shape[1]):
        table3.add_row([f+1, feat_labels[indices[f]], round(importances[indices[f]]*100, 2)])
        list.append(feat_labels[indices[f]])
    print(table3)
    prettytableprint('09_所有变量影响程度排序1', "9.所有变量影响程度排序")
    prettytableprint('09_所有变量影响程度排序2', table3)
    print()

    # 可视化
    import matplotlib.pyplot as plt
    plt.figure(figsize=(10, 6))
    plt.title("数据集中各个特征的重要程度", fontsize=18)
    plt.ylabel("import level", fontsize=15, rotation=90)
    plt.rcParams['font.sans-serif'] = ["SimHei"]
    plt.rcParams['axes.unicode_minus'] = False
    for i in range(x_columns.shape[0]):
        plt.bar(i, importances[indices[i]], color='blue', align='center')
        plt.xticks(np.arange(x_columns.shape[0]), list, rotation=90, fontsize=12)
    plt.savefig(tempt_path + '09_所有变量影响程度排序3' + '.png')
    plt.show()


# 画出决策树
def draw_tree():
    feat_labels = data_train.columns[1:len(data_train.columns) - 1]
    X_lables = [x for x in feat_labels]
    y_lables = data_train.columns[len(data_train.columns) - 1]

    # 接着，构建决策树模型
    model_tree = DecisionTreeClassifier(criterion='gini', max_depth=5, min_samples_split=2,min_samples_leaf =2,max_features=2)
    model_tree.fit(X_train, y_train)

    # # 评价模型准确性
    # y_prob = model_tree.predict_proba(X_test)[:, 1]
    # y_pred = np.where(y_prob > 0.5, 1, 0)
    # model_tree_score = model_tree.score(X_test, y_pred)
    # print('模型准确性：' + str(model_tree_score))

    os.environ["PATH"] += os.pathsep + 'C:/Program Files (x86)/Graphviz2.38/bin/'
    dot_tree = tree.export_graphviz(model_tree, out_file=None, feature_names=X_lables, class_names=y_lables,
                                    filled=True, rounded=True, special_characters=True)
    graph = pydotplus.graph_from_dot_data(dot_tree)
    img = Image(graph.create_png())
    # 设置决策树名称和输出路径
    prettytableprint("12_决策树1", "12.决策树")
    graph.write_png(tempt_path + "12_决策树2.png")
    print("成功输出决策树")


# 模型预测
def make_predictions():
    predictions = forest.predict(data_predict.drop(ID, axis=1))  # 删除ID字段
    output = pd.DataFrame({'id': data_predict[ID], 'prediction': predictions})
    output.to_csv(path + '预测结果.csv')
    output.head()
    print("Congratulation!  Finish all steps")


def write2word(tempt_path):
    doc = docx.Document()
    for x, y, files in os.walk(tempt_path):
        files.sort(key=lambda x: int(x[:2]))
        # print(files)

        for file in files:
            if file.split('.')[1] == 'txt':
                f = open(tempt_path + file, 'r', encoding='UTF-8')
                # print(f.read())
                # print(type(f.read()))
                doc.add_paragraph(f.read())
                # print('成功写入文字：' + file)
                f.close()
            elif file.split('.')[1] == 'png':
                doc.add_picture(tempt_path + file, width=Inches(5))
                # print('成功写入图片：' + file)
    doc.save(path + '训练报告.docx')
    shutil.rmtree(tempt_path)


if __name__ == '__main__':
    # 获取数据
    path = "F:/测试数据/机器学习/"
    tempt_path = path + 'tempt/'
    # 创建输出路径
    if os.path.isdir(tempt_path):
        shutil.rmtree(tempt_path)
    os.makedirs(tempt_path)     
        
    data_train = pd.read_csv(path + "traindata.csv")
    data_predict = pd.read_csv(path + "predictdata.csv")
    print(data_train.head())
    print('-----------------------------------------')
    print()
    print(data_train.describe())
    print()

    # 特征因子
    feature_column = [x for x in data_train.columns[1:len(data_train.columns) - 1]]
    # ID
    ID = data_train.columns[0]
    # 标签因子
    lable_column = data_train.columns[len(data_train.columns) - 1]
    print('--- Data summary ---', '特征因子：' + str(feature_column), '标签因子：' + lable_column, sep='\n')
    prettytableprint('01_特征因子', '1.特征因子：' + str(feature_column))
    prettytableprint('02_标签因子', '2.标签因子：' + lable_column)
    print()

    # 删除ID字段和标签字段，只留训练数据。
    X = data_train.drop([lable_column, ID], axis=1)
    y = data_train[lable_column]
    p = 0.3  # 验证数据占比（参数）
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=p, random_state=23)

    # 打印标签统计表
    print_featurelable()

    # 不同n_estimators取值对应误差大小
    findbest_n_estimators(1500)

    # 不同max_features取值对应误差大小
    findbest_maxfeatures()

    # 建立模型
    forest = build_model()

    # 单因子分析
    singlefeature(len(feature_column))

    # 打印混淆矩阵
    print_mxjuzhen(forest)

    # 交叉验证
    cross_validation()

    # 因子重要性判断
    top_variable_importance(forest)

    # 画出决策树
    draw_tree()

    # 模型预测
    make_predictions()

    # 将分析结果写入word
    write2word(tempt_path)





